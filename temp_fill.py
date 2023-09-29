def pictu(your_image_url):
    import requests
    from io import BytesIO
    response = requests.get(your_image_url)
    binary_img = BytesIO(response.content)
    return binary_img
def doc_fil(info):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt
    from datetime import date
    from docx import Document as doc
    from docx.shared import Inches
    from pathlib import Path
    from io import BytesIO
    temp_path = Path("C:\\Users\\novikov.rn\\Desktop\\VKOPatel\\vkop_template.docx")
    d = doc(temp_path)
    obj_styles = d.styles
    obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'
    outfile = BytesIO()
    d.tables[0].rows[1].cells[0].paragraphs[0].add_run(info[0], style = 'CommentsStyle').bold = True
    d.tables[1].rows[0].cells[1].text = info[1]
    d.tables[1].rows[1].cells[1].text = info[2]
    d.tables[1].rows[2].cells[1].text = date.today().strftime("%d.%m.%Y")
    d.tables[1].rows[0].cells[3].text = str(info[3])
    d.tables[1].rows[1].cells[3].text = str(info[4])
    d.tables[1].rows[2].cells[3].text = str(info[5])
    d.tables[2].rows[0].cells[0].paragraphs[0].add_run(str(info[6]), style = 'CommentsStyle').bold = True
    d.tables[3].rows[1].cells[1].text = str(info[7])+"м³/ч"
    d.tables[3].rows[2].cells[1].text = str(info[8])+"Па"
    d.tables[3].rows[2].cells[3].text = str(info[9])+"м³/ч"
    d.tables[3].rows[3].cells[3].text = str(int(info[10]))+"Па"
    d.tables[3].rows[2].cells[5].text = str(int(6000 / int(info[6].split("/")[1].split("-")[0]))) + "об/мин"
    d.tables[3].rows[5].cells[3].text = info[6].split("-")[-1]
    d.tables[3].rows[1].cells[5].text = str(int(info[6].split("/")[0].split("-")[-1])/100)+"кВт"
    d.tables[3].rows[4].cells[3].text = str(int(info[6].split("-")[1])*10)+"мм"
    d.tables[4].rows[0].cells[0].add_paragraph().add_run().add_picture(pictu(info[11]), width=Inches(4))
    d.tables[4].rows[0].cells[1].add_paragraph().add_run().add_picture(info[12], width=Inches(3.5))
    g = d.tables[5]
    gabs = [[key, info[-2][key][0]] for key in info[-2].keys()]
    for i in range(len(gabs)):
        for j in range(2):
            g.rows[j].cells[i].text = gabs[i][j]
    for k in range(len(info[-1])):
        d.tables[6].rows[k].cells[0].text = info[-1][k]
    d.save(outfile)
    return outfile




